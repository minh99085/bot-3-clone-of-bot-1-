"""Human-readable Word (.docx) performance report with section scores and history."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Optional, Union


def _fmt(v, nd=2):
    if v is None:
        return "—"
    if isinstance(v, float):
        return f"{v:.{nd}f}"
    return str(v)


def build_word_report(
    light: dict,
    *,
    status: Optional[dict] = None,
    ledger: Optional[dict] = None,
    score_history: Optional[dict] = None,
    output_path: Optional[Union[str, Path]] = None,
) -> bytes:
    """Build a Word report; return bytes and optionally write to ``output_path``."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    from engine.pulse.reporting import build_report_sections
    from engine.pulse.performance_scoring import compute_report_scores

    light = light or {}
    status = status or {}
    ledger = ledger or {}
    sections = light.get("sections") or build_report_sections(light, status=status, ledger=ledger)
    scores = light.get("scores") or compute_report_scores(
        sections, global_reconciled=bool(light.get("global_reconciled")))
    hist = score_history or light.get("score_history") or {}

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    title = doc.add_heading("BTC 5-Minute Pulse — Performance Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(
        f"PAPER ONLY  ·  reconciled={light.get('global_reconciled')}  ·  "
        f"ticks={status.get('ticks', '—')}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_heading("Performance Scorecard", level=1)
    overall = scores.get("overall", {})
    doc.add_paragraph(
        f"Overall score: {overall.get('score', '—')} / 100  (Grade: {overall.get('grade', '—')})")

    sc_table = doc.add_table(rows=4, cols=4)
    sc_table.style = "Table Grid"
    headers = ["Section", "Score", "Grade", "Weight"]
    for i, h in enumerate(headers):
        sc_table.rows[0].cells[i].text = h
    rows = [
        ("Trading Performance", scores.get("trading_performance", {}), "50%"),
        ("Operation", scores.get("operation", {}), "25%"),
        ("External Signals", scores.get("external_signals", {}), "25%"),
    ]
    for ri, (name, block, weight) in enumerate(rows, start=1):
        sc_table.rows[ri].cells[0].text = name
        sc_table.rows[ri].cells[1].text = _fmt((block or {}).get("score"), 1)
        sc_table.rows[ri].cells[2].text = str((block or {}).get("grade") or "—")
        sc_table.rows[ri].cells[3].text = weight

    doc.add_heading("Score History", level=2)
    entries = (hist.get("entries") or [])[-20:]
    if entries:
        ht = doc.add_table(rows=1 + len(entries), cols=7)
        ht.style = "Table Grid"
        hcols = ["UTC", "Settled", "Trading", "Operation", "Signals", "Overall", "Ticks"]
        for i, h in enumerate(hcols):
            ht.rows[0].cells[i].text = h
        for ri, e in enumerate(entries, start=1):
            sc = e.get("scores") or {}
            ht.rows[ri].cells[0].text = str(e.get("utc") or "")[:19]
            ht.rows[ri].cells[1].text = str(e.get("settled") or "")
            ht.rows[ri].cells[2].text = _fmt(sc.get("trading_performance"), 1)
            ht.rows[ri].cells[3].text = _fmt(sc.get("operation"), 1)
            ht.rows[ri].cells[4].text = _fmt(sc.get("external_signals"), 1)
            ht.rows[ri].cells[5].text = _fmt(sc.get("overall"), 1)
            ht.rows[ri].cells[6].text = str(e.get("ticks") or "")
    else:
        doc.add_paragraph("No score history yet — snapshots are recorded as the bot runs.")

    def _score_breakdown(score_block: dict):
        note = score_block.get("note")
        if note:
            p = doc.add_paragraph(note)
            p.runs[0].italic = True
        comps = score_block.get("components") or {}
        if comps:
            doc.add_heading("Score breakdown", level=2)
            ct = doc.add_table(rows=1 + len(comps), cols=3)
            ct.style = "Table Grid"
            ct.rows[0].cells[0].text = "Component"
            ct.rows[0].cells[1].text = "Score"
            ct.rows[0].cells[2].text = "Weight contribution"
            for ri, (name, info) in enumerate(sorted(comps.items()), start=1):
                ct.rows[ri].cells[0].text = name.replace("_", " ")
                ct.rows[ri].cells[1].text = _fmt((info or {}).get("score"), 1)
                ct.rows[ri].cells[2].text = _fmt((info or {}).get("contribution"), 1)

    tp = sections.get("trading_performance", {}) or {}
    hl = tp.get("headline", {}) or {}
    doc.add_page_break()
    doc.add_heading("1. Trading Performance — Detail", level=1)
    sb = scores.get("trading_performance", {})
    doc.add_paragraph(
        f"Section score: {sb.get('score', '—')} / 100  (Grade: {sb.get('grade', '—')})")
    _score_breakdown(sb)
    doc.add_heading("Key metrics", level=2)
    mt = doc.add_table(rows=10, cols=2)
    mt.style = "Table Grid"
    metrics = [
        ("Total on-hand", f"${_fmt(hl.get('total_on_hand_usd'))}"),
        ("Total return", f"{_fmt(hl.get('total_return_pct') or hl.get('return_pct'))}%"),
        ("Directional PnL", f"${_fmt(hl.get('directional_realized_pnl_usd'))}"),
        ("Win rate", _fmt(hl.get("win_rate"), 4)),
        ("Win rate UP / DOWN", f"{_fmt(hl.get('win_rate_up'), 4)} / {_fmt(hl.get('win_rate_down'), 4)}"),
        ("Profit factor", _fmt(hl.get("profit_factor"), 4)),
        ("Trades settled", f"{hl.get('trades')} / {hl.get('settled')}"),
        ("Avg PnL/trade", f"${_fmt((tp.get('ledger') or {}).get('avg_pnl_per_trade'))}"),
        ("Global reconciled", str((tp.get("reconciliation") or {}).get("global_reconciled"))),
    ]
    for i, (k, v) in enumerate(metrics):
        mt.rows[i].cells[0].text = k
        mt.rows[i].cells[1].text = v

    positions = tp.get("recent_positions") or []
    if positions:
        doc.add_heading("Recent positions (last 15)", level=2)
        pt = doc.add_table(rows=1 + min(len(positions), 15), cols=5)
        pt.style = "Table Grid"
        for i, h in enumerate(["Side", "Entry", "Outcome", "Won", "PnL"]):
            pt.rows[0].cells[i].text = h
        for ri, p in enumerate(positions[:15], start=1):
            pt.rows[ri].cells[0].text = str(p.get("side") or "")
            pt.rows[ri].cells[1].text = _fmt(p.get("entry_price"), 3)
            ou = p.get("outcome_up")
            pt.rows[ri].cells[2].text = "up" if ou else ("down" if ou is False else "—")
            won = p.get("won")
            pt.rows[ri].cells[3].text = "Yes" if won else ("No" if won is False else "—")
            pt.rows[ri].cells[4].text = f"${_fmt(p.get('pnl_usd'))}"

    op = sections.get("operation", {}) or {}
    doc.add_page_break()
    doc.add_heading("2. Operation — Detail", level=1)
    sb = scores.get("operation", {})
    doc.add_paragraph(
        f"Section score: {sb.get('score', '—')} / 100  (Grade: {sb.get('grade', '—')})")
    _score_breakdown(sb)
    eng = op.get("engine", {}) or {}
    doc.add_paragraph(
        f"Ticks: {eng.get('ticks')} · Reconciled: {eng.get('global_reconciled')} · "
        f"Readiness: {(op.get('readiness') or {}).get('status', '—')}")
    lc = op.get("candidate_lifecycle", {}) or {}
    doc.add_paragraph(f"Candidates created: {lc.get('created')} · Terminals: {lc.get('terminals')}")
    doc.add_paragraph(f"Rejected by stage: {lc.get('rejected_by_stage')}")
    loops = (op.get("loops", {}) or {}).get("loops", {}) or {}
    if loops:
        doc.add_heading("Sub-loops", level=2)
        lt = doc.add_table(rows=1 + len(loops), cols=4)
        lt.style = "Table Grid"
        lt.rows[0].cells[0].text = "Loop"
        lt.rows[0].cells[1].text = "Role"
        lt.rows[0].cells[2].text = "Trigger"
        lt.rows[0].cells[3].text = "Stop condition"
        for ri, (name, info) in enumerate(sorted(loops.items()), start=1):
            lt.rows[ri].cells[0].text = name
            lt.rows[ri].cells[1].text = str((info or {}).get("role") or "")
            lt.rows[ri].cells[2].text = str((info or {}).get("trigger") or "")
            lt.rows[ri].cells[3].text = str((info or {}).get("stop_condition") or "")[:60]

    ex = sections.get("external_signals", {}) or {}
    imp = ex.get("impact_summary", {}) or {}
    doc.add_page_break()
    doc.add_heading("3. External Signals — Detail", level=1)
    sb = scores.get("external_signals", {})
    doc.add_paragraph(
        f"Section score: {sb.get('score', '—')} / 100  (Grade: {sb.get('grade', '—')})")
    _score_breakdown(sb)
    doc.add_heading("Signal impact on trading", level=2)
    it = doc.add_table(rows=6, cols=2)
    it.style = "Table Grid"
    impact_rows = [
        ("TV aligned bot win rate", imp.get("tv_aligned_bot_win_rate")),
        ("TV opposed bot win rate", imp.get("tv_opposed_bot_win_rate")),
        ("TV signal hit rate", imp.get("tv_signal_hit_rate")),
        ("Grok direction accuracy", imp.get("grok_direction_accuracy")),
        ("Grok view accuracy", imp.get("grok_view_accuracy")),
        ("CEX-lead proven", imp.get("cex_lead_any_proven")),
    ]
    for i, (k, v) in enumerate(impact_rows):
        it.rows[i].cells[0].text = k
        it.rows[i].cells[1].text = _fmt(v, 4) if isinstance(v, float) else str(v)

    tv = ex.get("tradingview", {}) or {}
    doc.add_paragraph(
        f"TV alerts: received {tv.get('tradingview_alerts_received', 0)} · "
        f"valid {tv.get('tradingview_alerts_valid', 0)} · "
        f"rejected {tv.get('tradingview_alerts_rejected', 0)}")

    buf = BytesIO()
    doc.save(buf)
    data = buf.getvalue()
    if output_path is not None:
        Path(output_path).write_bytes(data)
    return data